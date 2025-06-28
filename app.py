
import streamlit as st
import pandas as pd
import numpy as np
import re
from io import BytesIO
import xlsxwriter
from docx import Document
from tempfile import NamedTemporaryFile

st.set_page_config(page_title="Autoevaluación & Plan de Carrera", layout="wide")

FILE_COMP = "competencias_agrupadas.xlsx"
FILE_BEH = "comportamientos_agrupados.xlsx"

@st.cache_data(show_spinner=True)
def load_data():
    df_comp = pd.read_excel(FILE_COMP)
    df_beh = pd.read_excel(FILE_BEH)
    return df_comp, df_beh

df_comp, df_beh = load_data()
competencias_cols = df_comp.columns[2:10].tolist()
col_comp = "Competencia"

@st.cache_data()
def build_behavior_dict(df_beh):
    behavior_dict = {}
    for _, row in df_beh.iterrows():
        job = row["Job Title"]
        area = row["Area"]
        comp = row[col_comp]
        beh = re.sub(r"^\d+\.\s*", "", str(row["Comportamientos"]).strip().lower())
        val = row["Valor"]
        key = (job, area)
        if pd.notna(job) and pd.notna(comp) and pd.notna(beh):
            behavior_dict.setdefault(key, {}).setdefault(comp, {})[beh] = val
    return behavior_dict

behavior_dict = build_behavior_dict(df_beh)

# UI
col1, col2 = st.columns([1, 8])
with col1:
    st.image("https://cdn-icons-png.flaticon.com/512/3135/3135767.png", width=80)
with col2:
    st.title("Autoevaluación de Competencias y Comportamientos")

nombre = st.text_input("Nombre completo")
areas_unique = sorted(df_comp["Area"].dropna().unique())
area_sel = st.selectbox("Área", ["-- Selecciona --"] + areas_unique)
puestos_sel = sorted(df_comp[df_comp["Area"] == area_sel]["Job Title"].unique()) if area_sel != "-- Selecciona --" else []
puesto = st.selectbox("Puesto actual", ["-- Selecciona --"] + puestos_sel)

st.header("1️⃣ Reparte 100 puntos entre las 8 competencias")
cols = st.columns(4)
comp_values = {comp: st.number_input(comp, 0, 100, 0, 1, key=f"comp_{i}") for i, comp in enumerate(competencias_cols)}
total_comp = sum(comp_values.values())
st.markdown(f"**Total asignado:** {total_comp} / 100")

st.header("2️⃣ Evalúa los comportamientos (1‑5)")
beh_values = {}
for comp in competencias_cols:
    st.subheader(comp)
    filtered = df_beh[df_beh[col_comp] == comp]["Comportamientos"].dropna().unique()
    for beh in sorted(filtered):
        clean = re.sub(r"^\d+\.\s*", "", beh.strip())
        beh_values[clean.lower()] = st.slider(clean, 1, 5, 3, key=f"beh_{comp}_{clean}")

if st.button("Generar plan de carrera"):
    if area_sel == "-- Selecciona --" or puesto == "-- Selecciona --":
        st.error("Selecciona tu área y puesto actual.")
        st.stop()
    if total_comp != 100:
        st.error("Distribuye exactamente 100 puntos entre las competencias.")
        st.stop()
    if not nombre:
        st.error("Por favor, introduce tu nombre.")
        st.stop()

    ipe_actual = df_comp.loc[(df_comp["Job Title"] == puesto) & (df_comp["Area"] == area_sel), "IPE_val"].iloc[0]
    df_persona = pd.Series(comp_values)
    pesos = df_persona / 100
    resultados = []

    for _, row in df_comp.iterrows():
        if pd.isna(row["IPE_val"]) or row["IPE_val"] < ipe_actual:
            continue

        gap_comp = (abs(df_persona - row[competencias_cols]) * pesos).sum()
        job_beh = behavior_dict.get((row["Job Title"], row["Area"]), {})
        gap_beh, total_peso = 0, 0

        for comp in competencias_cols:
            comp_beh = job_beh.get(comp, {})
            peso_comp = pesos[comp]
            for beh, val in comp_beh.items():
                if beh in beh_values:
                    gap = abs(beh_values[beh] - val)
                    gap_beh += gap * peso_comp
                    total_peso += peso_comp

        gap_beh = gap_beh / total_peso if total_peso else 0
        gap_total = 0.7 * gap_comp + 0.3 * gap_beh

        resultados.append({
            "Job Title": row["Job Title"],
            "Area": row["Area"],
            "IPE": row["IPE_val"],
            "Gap Total": round(gap_total, 2),
            "Gap Comp": round(gap_comp, 2),
            "Gap Beh": round(gap_beh, 2)
        })

    df_r = pd.DataFrame(resultados).sort_values(["IPE", "Gap Total"]).reset_index(drop=True)
    resumen = df_r.drop_duplicates(subset=["Job Title", "Area"])
    st.subheader("🔍 Plan de carrera personalizado")
    st.dataframe(resumen, use_container_width=True)

    doc = Document()
    doc.add_heading(f"Plan de Carrera: {nombre}", 0)
    doc.add_paragraph(f"Puesto actual: {puesto} ({area_sel}) - IPE {ipe_actual}")
    doc.add_paragraph("\nFortalezas principales:")
    for c in df_persona.sort_values(ascending=False).head(2).index:
        doc.add_paragraph(f"- {c} ({df_persona[c]} puntos)", style='ListBullet')

    doc.add_paragraph("\nDesarrollo recomendado:")
    for _, r in resumen.iterrows():
        doc.add_heading(r["Job Title"], level=2)
        doc.add_paragraph(f"Área: {r['Area']} | IPE: {r['IPE']} | Gap Total: {r['Gap Total']}")

    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        resumen.to_excel(writer, index=False, sheet_name="Resumen")

    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp.seek(0)
        doc_bytes = tmp.read()

    st.download_button("📥 Descargar Word del Plan", data=doc_bytes, file_name=f"plan_carrera_{nombre.replace(' ', '_')}.docx")
    st.download_button("📥 Descargar Excel del Plan", data=buffer.getvalue(), file_name=f"plan_carrera_{nombre.replace(' ', '_')}.xlsx")
